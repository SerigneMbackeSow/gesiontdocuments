####################Test Archive
from datetime import datetime, date
from io import BytesIO

from django.core.exceptions import ObjectDoesNotExist
from django.db.models import Q
from django.http import JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt

from archivedocuments.models import Utilisateur, Dossiers, Documents, Acces, Demandes, Direction

############Import Dossier

@csrf_exempt
def testarchive(request, id_user):
    user = Utilisateur.objects.get(id_user=id_user)
    return render(request, 'templatetra/testarchive.html', { 'user': user, 'util': user})
#################Gestion Dossier
def liste_dossier(request, id_user, id_prt):
    # Obtenir le dossier courant
    dossier_crt = Dossiers.objects.get(id_dossier=id_prt)
    user = Utilisateur.objects.get(id_user=id_user)
    dossier_racine = Dossiers.objects.get(id_parent__isnull=True)

    # Fonction récursive pliste_dosour obtenir tous les sous-dossiers et documents
    def get_descendants(dossier):
        sous_dossiers = Dossiers.objects.filter(id_parent=dossier.id_dossier, direction=user.direction)
        documents = Documents.objects.filter(id_parent=dossier.id_dossier)
        all_sous_dossiers = list(sous_dossiers)
        all_documents = list(documents)

        for sous_dossier in sous_dossiers:
            desc_docs, desc_dossiers = get_descendants(sous_dossier)
            all_documents.extend(desc_docs)
            all_sous_dossiers.extend(desc_dossiers)

        return all_documents, all_sous_dossiers

    lst_document, lst_dossier = get_descendants(dossier_crt)

    # Filtrer pour ne garder que les documents ayant pour parent dossier_crt
    lst_document = [doc for doc in lst_document if doc.id_parent == dossier_crt.id_dossier]

    # Mettre à jour les accès et états pour les documents et dossiers
    def update_acces_and_etat(items, is_document=True):
        for item in items:
            try:
                acces = Acces.objects.get(id_utilisateur=id_user, etat=1, permission='acces',
                                          id_objet=(item.id_document if is_document else item.id_dossier))
                item.acces = acces.etat
            except Acces.DoesNotExist:
                item.acces = 5

            try:
                dmd = Demandes.objects.get(id_demandeur=id_user,
                                           id_objet=(item.id_document if is_document else item.id_dossier))
                item.etat = dmd.etat
            except Demandes.DoesNotExist:
                item.etat = 5

    update_acces_and_etat(lst_document)
    update_acces_and_etat(lst_dossier, is_document=False)

    # Construire l'arborescence des dossiers parents
    arborescence = []
    get_parent = dossier_crt
    while get_parent:
        arborescence.append(get_parent)
        get_parent = Dossiers.objects.filter(id_dossier=get_parent.id_parent).first()

    # Gestion de tous les documents
    tout_document = Documents.objects.all()
    for doc in tout_document:
        dossier_acces = Acces.objects.filter(etat=1, id_utilisateur=id_user, permission='acces', type='dossier')
        restricted_doc_ids = set(dossier_acces.values_list('id_objet', flat=True))
        doc.acces_dos = 1 if doc.id_parent in restricted_doc_ids else 0

    update_acces_and_etat(tout_document)

    return render(request, 'templatetra/liste_dossier.html', {
        'util': user,
        'dossier_racine': dossier_racine,
        #'tout_doc': tout_document,
        #'tout_dos': Dossiers.objects.all(),
        'dossier_crt': dossier_crt,
        #'lst_doc': lst_document,
        #'lst_dos': lst_dossier,
        'arborescence': reversed(arborescence),
    })


def dossier_racine(request, id_user):
    user = Utilisateur.objects.get(id_user=id_user)
    dossier_crt = None
    id_racine = None

    try:
        dossier_crt = Dossiers.objects.get(id_parent__isnull=True)
        id_racine = dossier_crt.id_dossier
    except Dossiers.DoesNotExist:
        pass
    if user.direction == 'management' :
        return redirect(f'/archive/liste_dossier/{id_user}/{id_racine}')

    lst_dossier = Dossiers.objects.filter(id_parent=id_racine,direction=user.direction)
    arborescence=[]
    arborescence.append(dossier_crt)
#    dos_tom=Dossiers.objects.get(nom='TOM')
   # dossier_crt=Dossiers.objects.get(id_parent=dos_tom.id_dossier,direction=user.direction)


    return render(request, 'templatetra/liste_dossier_raccine.html', {
        'util': user,
        'lst_dos': lst_dossier,
        'dossier_crt': dossier_crt,
        'arborescence': arborescence
    })


from django.shortcuts import redirect, render
from .models import Dossiers
from .forms import DossiersForm


def ajouter_dossier(request):
    if request.method == 'POST':
        direction = request.POST.get('direction')
        nom = request.POST.get('nom')
        id_parent = request.POST.get('id_parent')
        id_user = request.POST.get('id_user')
       # if not nom:
            # Vous pouvez rediriger vers une page d'erreur ou afficher un message si `nom` est vide
        #    return render(request, 'templatetra/ajouter_dossier.html', {'error': 'Le nom est requis.'})
        user = Utilisateur.objects.get(id_user=id_user)
        if user.direction == 'ADMIN':
            path = f'/archive/liste_dossier_admin/{id_user}/{id_parent}'
        else:
            path = f'/archive/liste_dossier/{id_user}/{id_parent}'
        # Créez le dossier
        Dossiers.objects.create(
            direction=direction,
            nom=nom,
            date_creation=datetime.now(),
            id_parent=id_parent or None
        )

        return redirect(path)

@csrf_exempt
def ajouterdocument(request):
    if request.method == 'POST':
        user = None
        dossier = None
        message = ''
        messagebl = ''

        try:
            user = Utilisateur.objects.get(id_user=request.POST['id_user'])
            dossier = Dossiers.objects.get(id_dossier=request.POST['id_parent'])
        except Utilisateur.DoesNotExist:
            message = "Utilisateur introuvable."
        except Dossiers.DoesNotExist:
            message = "Dossier introuvable."

        if not user or not dossier:
            return redirect(f"/archive/liste_dossier/{request.POST['id_user']}/{request.POST['id_parent']}")
            #return render(request, 'templatetra/ajout-document.html', {'message': message})

        fichier = request.FILES.get('fichier')
        if fichier:
            handle_uploaded_file(fichier)
        else:
            return redirect(f"/archive/liste_dossier/{request.POST['id_user']}/{request.POST['id_parent']}")
            #return render(request, 'templatetra/ajout-document.html', {'message': "Fichier manquant."})

        current_timestamp = int(datetime.timestamp(datetime.now()))

        if Documents.objects.filter(numero_document=request.POST['numero']).exists():
            message = "L'identifiant mentionné existe déjà."

        if Documents.objects.filter(bl=request.POST['bl']).exists():
            messagebl = "Le BL mentionné existe déjà."

        if message or messagebl:
            return redirect(f"/archive/liste_dossier/{request.POST['id_user']}/{request.POST['id_parent']}")
            #return render(request, 'templatetra/ajout-document.html', {'message': message, 'messagebl': messagebl})

        instance = Documents(
            numero_document=request.POST['numero'],
            date_creation=datetime.now(),
            eta=request.POST.get('eta') or '1900-01-01T00:00',
            bl=request.POST.get('bl'),
            client=request.POST.get('client'),
            vessel=request.POST.get('navire'),
            numero_voyage=request.POST.get('voyage'),
            id_parent=request.POST.get('id_parent'),
            journal=request.POST.get('journal'),
            nom=str(current_timestamp) + fichier.name,
            montant=request.POST.get('montant') or 0,
            conteneur=request.POST.get('conteneur'),
            direction=user.direction
        )
        instance.save()

        return redirect(f"/archive/liste_dossier/{request.POST['id_user']}/{request.POST['id_parent']}")

    else:
        id_user = request.GET.get('id_user')
        dossier_id = request.GET.get('dossiers')
        try:
            dossier = Dossiers.objects.get(id_boite=dossier_id)
            user = Utilisateur.objects.get(id_user=id_user)
        except (Dossiers.DoesNotExist, Utilisateur.DoesNotExist):
            return render(request, 'docs/ajouterdossier.html', {'message': "Erreur de chargement des données."})

        return render(request, 'docs/ajouterdossier.html', {'user': user, 'dossier': dossier})
@csrf_exempt
def handle_uploaded_file(f):
    current_timestamp = int(datetime.timestamp(datetime.now()))
    with open('static/archives/' + str(current_timestamp) + f.name, 'wb+') as destination:
        for chunk in f.chunks():
            destination.write(chunk)
from .models import Documents, Dossiers, Utilisateur
def convert_to_pdf(file_path):
    output_file = os.path.splitext(file_path)[0] + '.pdf'
    command = ['libreoffice', '--headless', '--convert-to', 'pdf', file_path]
    subprocess.run(command, check=True)
    return output_file
def affiche_document(request, id_user, id_doc):
    document = Documents.objects.get(id_document=id_doc)
    dossier_crt = Dossiers.objects.get(id_dossier=document.id_parent)
    user = Utilisateur.objects.get(id_user=id_user)
    base_dir = r'C:\MEDLOG\gestion_archive\inventory-management\static\archives'
    file_path = os.path.join(base_dir, document.nom)
    file_extension = os.path.splitext(file_path)[1].lower()
    try:
        if file_extension == '.pdf':
            return render(request, 'templatetra/affiche_document.html', {
                'util': user,
                #'aff_tt':'aff',
                'doc': document,
                'dossier_crt': dossier_crt,
                'file_url': f"/static/archives/{document.nom}",
                'file_type': 'pdf'
            })
        elif file_extension in ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx']:
            output_file = convert_to_pdf(file_path)
            return render(request, 'templatetra/affiche_document.html', {
                'util': user,
                'doc': document,
                'dossier_crt': dossier_crt,
                'file_url': f"/static/archives/{os.path.basename(output_file)}",
                'file_type': 'pdf'
            })
        else:
            return HttpResponse("Type de fichier non pris en charge.", status=415)
    except Exception as e:
        response = HttpResponse(open(file_path, 'rb').read(), content_type='application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
        return response
import subprocess
def convert_to_pdf(file_path):
    output_file = os.path.splitext(file_path)[0] + '.pdf'
    command = ['libreoffice', '--headless', '--convert-to', 'pdf', file_path]
    subprocess.run(command, check=True)
    return output_file


########################### RECHERCHE GLOBALE ###########################################
####################Gestion des acces
def ajout_acces(request) :
    if request.method == 'POST':
        if request.POST['tout_doc'] == 'tout_doc' :
            user = request.POST['id_util']
            doscrt = request.POST.get('doscrt')
            instance = Acces(
                permission='acces',
                etat=1,
                type=request.POST['type'],
                id_objet=request.POST['id_objet'],
                id_utilisateur=request.POST['id_user'],
            )
            instance.save()
            # return  redirect(f'/archive/liste_dossier/{user}/{doscrt}')
            return redirect(f'/archive/liste_document_direction/{user}')
        else :
            doscrt = request.POST.get('doscrt')

            user= request.POST['id_util']
            userdir=Utilisateur.objects.get(id_user=request.POST['id_user'])
            docuacc=Documents.objects.get(id_document=request.POST['id_objet'])
            if userdir.direction == docuacc.direction :
                doscrt = request.POST.get('doscrt')
                instance=Acces(
                    permission='acces',
                    etat=1,
                    type=request.POST['type'],
                    id_objet= request.POST['id_objet'],
                    id_utilisateur=request.POST['id_user'],
                )
                instance.save()
                #return  redirect(f'/archive/liste_dossier/{user}/{doscrt}')
                return redirect(f'/archive/liste_dossier/{user}/{doscrt}')
            else :
                acces=Acces.objects.get(type='document',etat=0,id_objet=docuacc.id_document,id_utilisateur=request.POST['id_user'])
                dmd = Demandes.objects.get(type='document', etat=1, id_objet=docuacc.id_document,
                                          id_demandeur=request.POST['id_user'])
                #acces.etat=10
                acces.delete()
                dmd.delete()

                return redirect(f'/archive/liste_dossier/{user}/{doscrt}')


def restraindre_page(request, id_user,id_objet,type):
    dossier_crt=False
    user = Utilisateur.objects.get(id_user=id_user)

    #user_direction = Utilisateur.objects.filter(direction=user.direction).exclude(user)
    acces_users = Acces.objects.filter(etat=1, id_objet=id_objet,type=type).values_list('id_utilisateur', flat=True)

  #####Agent direction
    agents_dir = Utilisateur.objects.filter(direction=user.direction).exclude(id_user__in=acces_users).exclude(id_user=user.id_user)
    ###################Agent No direction
    acces_users_no_dir = Acces.objects.filter(etat=0, id_objet=id_objet, type=type).values_list('id_utilisateur', flat=True)
    agents_no_dir = Utilisateur.objects.filter(id_user__in=acces_users_no_dir).exclude(direction=user.direction)
    if type == 'dossier' :
        objet=Dossiers.objects.get(id_dossier=id_objet)
        #id_objet = objet.id_dossier
        dossier_crt = Dossiers.objects.filter(id_dossier=objet.id_parent).first()
        if not dossier_crt :
            dossier_crt = Dossiers.objects.filter(id_dossier=objet.id_dossier).first()
        id_dos_crt = dossier_crt.id_dossier
        type='dossier'
    else :
       objet= Documents.objects.get(id_document=id_objet)
       #id_objet=objet.id_document
       dossier_crt=Dossiers.objects.get(id_dossier=objet.id_parent)
       id_dos_crt=dossier_crt.id_dossier
       type='document'

    return render(request, 'templatetra/restreindre_acces.html', {
        'util': user,
        'agents_no_dir' :agents_no_dir,
        'agents_dir': agents_dir,
        'type' :type,
        'id_objet' :id_objet,
        'id_dos_crt': id_dos_crt  # Pour avoir l'ordre du plus ancien au plus récent
    })


############gestion des deamndes
def ajout_demande(request) :
    if request.method == 'POST':

        # Récupérer les données du formulaire
        dos_crt=False
        iddoscrt = request.POST.get('dossier')
        type_dmd = request.POST.get('type')
        commentaire = request.POST.get('commentaire')
        id_objet = request.POST.get('objet')
        if request.POST.get('tout_doc') :

            user=Utilisateur.objects.get(id_user=request.POST.get('id_user'))
            if type_dmd == 'dossier' :
                dos=Dossiers.objects.get(id_dossier=id_objet)
                dos_crt=dos.id_parent
                recepteur=Utilisateur.objects.get(direction=dos.direction,role='chef')
            else :
                docu=Documents.objects.get(id_document=id_objet)
                dos = Dossiers.objects.get(id_dossier=docu.id_parent)
                recepteur = Utilisateur.objects.get(direction=dos.direction, role='chef')
                dos_crt=docu.id_parent
            dmd_exit=False
            try :
              dmd_exit=Demandes.objects.get(type=type_dmd,id_demandeur=user.id_user,id_objet=id_objet)
            except :
                pass
            if dmd_exit :
                dmd_exit.etat=0
                dmd_exit.save()
            else :

                nouvelle_demande = Demandes(
                    type=type_dmd,
                    commentaire=commentaire,
                    id_demandeur=user.id_user,
                    id_recepteur=recepteur.id_user,
                    etat=0,
                    id_objet=id_objet
                )
                nouvelle_demande.save()
            #return  redirect(f'/archive/liste_dossier/{user}/{doscrt}')
            return redirect(f'/archive/liste_document_direction/{user.id_user}')
        else :
            user = Utilisateur.objects.get(id_user=request.POST.get('id_user'))
            if type_dmd == 'dossier':
                dos = Dossiers.objects.get(id_dossier=id_objet)
                dos_crt = dos.id_parent
                recepteur = Utilisateur.objects.get(direction=dos.direction, role='chef')
            else:
                docu = Documents.objects.get(id_document=id_objet)
                dos = Dossiers.objects.get(id_dossier=docu.id_parent)
                recepteur = Utilisateur.objects.get(direction=dos.direction, role='chef')
                dos_crt = docu.id_parent
            dmd_exit = False
            try:
                dmd_exit = Demandes.objects.get(type=type_dmd, id_demandeur=user.id_user, id_objet=id_objet)
            except:
                pass
            if dmd_exit:
                dmd_exit.etat = 0
                dmd_exit.save()
            else:

                nouvelle_demande = Demandes(
                    type=type_dmd,
                    commentaire=commentaire,
                    id_demandeur=user.id_user,
                    id_recepteur=recepteur.id_user,
                    etat=0,
                    id_objet=id_objet
                )
                nouvelle_demande.save()
                docu=Documents.objects.get(id_document=id_objet)
                if docu.direction != user.direction:
                    instance = Acces(
                        permission='acces',
                        etat=1,
                        type='document',
                        id_objet=id_objet,
                        id_utilisateur=user.id_user
                    )
                    instance.save()
            # return  redirect(f'/archive/liste_dossier/{user}/{doscrt}')
            return redirect(f'/archive/liste_dossier/{user.id_user}/{iddoscrt}')
#################Gestion demande
def demande_permission_recus(request,id_user) :
    user=Utilisateur.objects.get(id_user=id_user)
    dmd= Demandes.objects.filter(etat=0,id_recepteur=id_user)
    return render(request, 'templatetra/demande_acces_recus.html', {'util': user,'lst_dmd': dmd})

def fetch_demande_acces(request, id_user):
    id_user = request.GET.get('id_user')
    if id_user:
        try:
            user = Utilisateur.objects.get(id_user=id_user)
            dmd = Demandes.objects.filter(etat=0, id_recepteur=id_user)
            data = []
            for doc in dmd:
                demandeur = Utilisateur.objects.get(id_user=doc.id_demandeur)
                data.append({
                    'id': doc.id_dmd,
                    'type': doc.type,
                    'direction': demandeur.direction,
                    'fullname': demandeur.fullname
                })
            return JsonResponse({'data': data})
        except ObjectDoesNotExist:
            return JsonResponse({'error': 'User not found'}, status=404)
    else:
        return JsonResponse({'error': 'id_user parameter is missing'}, status=400)
def accepter_dmd(request,id_dmd,id_user,) :
    dmd=Demandes.objects.get(id_dmd=id_dmd)
    dmd.commentaire_reponse='acceptée'
    dmd.etat=1
    demander=dmd.id_demandeur
    id_objet=dmd.id_objet
    if dmd.type == 'dossier' :
        acces=Acces.objects.get(type='dossier',id_utilisateur=demander,id_objet=id_objet,etat=1)
        acces.etat=0
        acces.save()
    if dmd.type == 'document' :
        acces=Acces.objects.get(type='document',id_utilisateur=demander,id_objet=id_objet,etat=1)
        acces.etat=0
        acces.save()
    dmd.save()
    user=Utilisateur.objects.get(id_user=id_user)
    return  redirect('/archive/demande_permission_recue/'+ str(user.id_user))
def refuser_dmd(request) :
    id_dmd=request.POST.get('id_dmd')
    id_user= request.POST.get('id_user')
    comentaire = request.POST.get('commentaire')
    dmd=Demandes.objects.get(id_dmd=id_dmd)
    dmd.etat=2
    demander = dmd.id_demandeur
    id_objet = dmd.id_objet
    dmd.commentaire_reponse=comentaire
    if dmd.type == 'document' :
        acces=Acces.objects.get(type='document',id_utilisateur=demander,id_objet=id_objet,etat=1)
        acces.etat=1
        acces.save()
    if dmd.type == 'dossier':
            acces = Acces.objects.get(type='dossier', id_utilisateur=demander, id_objet=id_objet,etat=1)
            acces.etat = 1
            acces.save()
    dmd.save()
    user=Utilisateur.objects.get(id_user=id_user)
    return  redirect('/archive/demande_permission_recue/'+ str(user.id_user))
@csrf_exempt
def demande_permission_envoyes(request, id_user):
    user = Utilisateur.objects.get(id_user=id_user)
    user = Utilisateur.objects.get(id_user=id_user)
    dmd = Demandes.objects.filter(etat=0, id_demandeur=id_user)
    return render(request, 'templatetra/demande_acces_envoyes.html', {'util': user, 'lst_dmd': dmd})


def fetch_demande_acces_envoyes(request, id_user):
    id_user = request.GET.get('id_user')
    if id_user:
        try:
            user = Utilisateur.objects.get(id_user=id_user)
            dmd = Demandes.objects.filter(id_demandeur=id_user)
            data = []
            for doc in dmd:
                demandeur = Utilisateur.objects.get(id_user=doc.id_demandeur)
                data.append({
                    'id': doc.id_dmd,
                    'type': doc.type,
                    'direction': demandeur.direction,
                    'fullname': demandeur.fullname,
                    'date_dmd':doc.date_dmd,
                    'etat': doc.etat,
                'commentaire': doc.commentaire_reponse
                })
            return JsonResponse({'data': data})
        except ObjectDoesNotExist:
            return JsonResponse({'error': 'User not found'}, status=404)
    else:
        return JsonResponse({'error': 'id_user parameter is missing'}, status=400)


def fetch_arborescence(request, id_docrt):
    id_doscrt = request.GET.get('id_dos')
    docu=Documents.objects.get(id_document=id_doscrt)
    dossier_crt = Dossiers.objects.get(id_dossier=docu.id_parent)
    arborescence = []
    get_parent = dossier_crt
    # Construire l'arborescence des dossiers parents
    while get_parent:
        arborescence.append(get_parent.nom)  # Remplacez 'nom' par l'attribut souhaité
        # Récupérer l'objet parent du dossier actuel
        get_parent = Dossiers.objects.filter(id_dossier=get_parent.id_parent).first()
    return JsonResponse({'data': arborescence})

##################Gestion Utilisateur
def gestion_utilsateurs(request,id_user):
    user = Utilisateur.objects.get(id_user=id_user)
    tout_user = Utilisateur.objects.all().exclude(direction='ADMIN')
    all_direction=Direction.objects.all()
    #all_direction = set(all_user.values_list('direction', flat=True))
    #direction_null = Dossiers.objects.all().exclude(direction__in=all_direction)
    return render(request, 'templatetra/gestion_utilisateur_archive.html', {'util': user, 'tout_user': tout_user,'lst_dir' :all_direction})

def ajouter_utilisateur(request):
    if request.method == 'POST':
        fullname = request.POST.get('fullname')
        email = request.POST.get('email')
        password = request.POST.get('password')
        direction = request.POST.get('direction')
        role = request.POST.get('role')
        if role == 'chef' :
            if  Utilisateur.objects.filter(direction=request.POST['direction'],role='chef').exists():
                user = Utilisateur.objects.get(direction='ADMIN')
                return redirect('/archive/gestion_utilsateurs/' + str(user.id_user))
        if Utilisateur.objects.filter(email=request.POST['email']).exists():
            user = Utilisateur.objects.get(direction='ADMIN')
            return redirect('/archive/gestion_utilsateurs/' + str(user.id_user))
        # Créez une nouvelle instance d'utilisateur
        nouvel_utilisateur = Utilisateur(
            fullname=fullname,
            email=email,
            password=password,  # Assurez-vous de hacher le mot de passe avant de le sauvegarder
            direction=direction,
            role=role
        )
        nouvel_utilisateur.save()  # Enregistrez l'utilisateur dans la base de données

        # Redirigez vers une page de succès ou retournez un message de succès
        user=Utilisateur.objects.get(direction='ADMIN')
        return redirect('/archive/gestion_utilsateurs/' + str(user.id_user))
    return render(request, 'ajouter_utilisateur.html')
###Actuver/ Desactiver
from django.shortcuts import redirect
from django.contrib import messages
def activer_utilisateur(request, id_user):
    try:
        user = Utilisateur.objects.get(id_user=id_user)
        user.status = 'actif'
        user.save()
        messages.success(request, 'Utilisateur activé avec succès!')
    except Utilisateur.DoesNotExist:
        messages.error(request, 'Utilisateur non trouvé!')
    except Exception as e:
        messages.error(request, f'Une erreur est survenue: {str(e)}')

    return redirect('/archive/gestion_utilsateurs/' + str(id_user))


def desactiver_utilisateur(request, id_user):
    try:
        user = Utilisateur.objects.get(id_user=id_user)
        user.status = 'innactif'
        user.save()
        messages.success(request, 'Utilisateur désactivé avec succès!')
    except Utilisateur.DoesNotExist:
        messages.error(request, 'Utilisateur non trouvé!')
    except Exception as e:
        messages.error(request, f'Une erreur est survenue: {str(e)}')

    return redirect('/archive/gestion_utilsateurs/' + str(id_user))



def reinitialiser_utilisateur(request, id_user):
    try:
        user = Utilisateur.objects.get(id_user=id_user)
        user.password = 'reinit'
        user.save()
        messages.success(request, 'Utilisateur activé avec succès!')
    except Utilisateur.DoesNotExist:
        messages.error(request, 'Utilisateur non trouvé!')
    except Exception as e:
        messages.error(request, f'Une erreur est survenue: {str(e)}')

    return redirect('/archive/gestion_utilsateurs/' + str(id_user))
@csrf_exempt
def update_utilisateur(request):
    agt = Utilisateur.objects.get(id_user=request.POST['id_user'])
    # Récupérer l'utilisateur qui effectue la mise à jour
    #user = Utilisateur.objects.get(id_user=request.POST['id_user'])
    # Mettre à jour les champs selon les données reçues en POST
    if request.POST.get('fullname'):
        agt.fullname= request.POST['fullname']

    if request.POST.get('email'):
        # Vérifier si l'email est déjà utilisé par un autre utilisateur
        if agt.email != request.POST['email'] and Utilisateur.objects.filter(email=request.POST['email']).exclude(
                id_user=agt.id_user).exists():
            message = 'Cet email est déjà utilisé par un autre utilisateur.'
            userad=Utilisateur.objects.get(direction='ADMIN')
            return redirect('/archive/gestion_utilsateurs/' + str(userad.id_user))

        else:
            agt.email = request.POST['email']

    if request.POST.get('role'):
        if Utilisateur.objects.filter(direction=agt.direction,role='chef').exclude(id_user=agt.id_user).exists():
            userad = Utilisateur.objects.get(direction='ADMIN')
            return redirect('/archive/gestion_utilsateurs/' + str(userad.id_user))
        agt.role = request.POST['role']
    if request.POST.get('direction'):
        agt.direction = request.POST['direction']
    if request.POST.get('role'):
        agt.role = request.POST['role']

    # Sauvegarder les modifications dans la base de données
    agt.save()
    userad = Utilisateur.objects.get(direction='ADMIN')
    return redirect('/archive/gestion_utilsateurs/' + str(userad.id_user))


#############Gestion tout Dossier
def tout_dossier_admin(request, id_user):
    user = Utilisateur.objects.get(id_user=id_user)
    tout_dossier = Dossiers.objects.get(id_parent__isnull=True)
    id_racine = tout_dossier.id_dossier
    direction_null = Dossiers.objects.filter(id_parent=id_racine)
    restricted_doc_ids = set(direction_null.values_list('direction', flat=True))
    lst_dir = Direction.objects.all().exclude(nom__in=restricted_doc_ids)
    arborescence = [Dossiers.objects.filter(id_dossier__isnull=True).first()]
    return redirect(f'/archive/liste_dossier_admin/{user.id_user}/{tout_dossier.id_dossier}')
    #return render(request, 'templatetra/tout_dossier_admin.html', {
        #'ldos': tout_dossier,
       # 'lst_dir': lst_dir,
       # 'util': user,
       # 'arborescence': arborescence,
    #})


def liste_dossier_admin(request, id_user, id_dos):
    tout_dossier = Dossiers.objects.filter(id_parent=id_dos)
    dos_crt = Dossiers.objects.get(id_dossier=id_dos)
    user = Utilisateur.objects.get(id_user=id_user)
    lst_dir = Direction.objects.all()

    arborescence = []
    get_parent = dos_crt
    # Construire l'arborescence des dossiers parents
    while get_parent:
        arborescence.append(get_parent)
        # Récupérer l'objet parent du dossier actuel
        get_parent = Dossiers.objects.filter(id_dossier=get_parent.id_parent).first()

    # Inverser l'arborescence pour avoir du plus ancien au plus récent
    arborescence.reverse()
    tout_dossier_dir = Dossiers.objects.get(id_parent__isnull=True)
    direction_null = Dossiers.objects.filter(id_parent=tout_dossier_dir.id_dossier)
    restricted_doc_ids = set(direction_null.values_list('direction', flat=True))
    lst_dir = Direction.objects.all().exclude(nom__in=restricted_doc_ids)
    return render(request, 'templatetra/liste_dossier_admin.html', {
        'tout_dos': tout_dossier,
        'arborescence': arborescence,
        'lst_dir': lst_dir,
        'dos_crt': dos_crt,
        'util': user,
    })



#########Modification Dossier

def modifier_dossier(request):
    if request.method == 'POST':
        user = Utilisateur.objects.get(id_user=request.POST.get('id_user'))
        if request.POST.get('id_dossier'):
            id_dossier = request.POST.get('id_dossier')
            dossier = Dossiers.objects.get(id_dossier=id_dossier)
            nom = request.POST.get('nom')
            if Dossiers.objects.filter(id_parent=dossier.id_parent,nom=nom).exclude(id_dossier=id_dossier).exists():
                return redirect("/archive/liste_dossier_admin/" + str(user.id_user) + "/" + str(dossier.id_parent))
            else :
                dossier.nom = nom
                dossier.save()
                if dossier.id_parent :
                 return redirect("/archive/liste_dossier_admin/" + str(user.id_user) + "/" + str(dossier.id_parent))
                else :
                    return redirect("/archive/tout_dossier_admin/" + str(user.id_user))
        else :
            id_dossier = request.POST.get('id_dossier')
            nouv_direction = request.POST.get('nouv_direction')
            nom = request.POST.get('nom')
            dossier=Dossiers.objects.get(id_dossier=id_dossier)
            dossier.direction=nouv_direction
            dossier.nom=nom
            dossier.save()
            return redirect("/archive/tout_dossier_admin/" + str(user.id_user))

def modifier_dossier_direction(request):
    if request.method == 'POST':
        user = Utilisateur.objects.get(id_user=request.POST.get('id_user'))
        if request.POST.get('id_dossier'):
            id_dossier = request.POST.get('id_dossier')
            dossier = Dossiers.objects.get(id_dossier=id_dossier)
            nom = request.POST.get('nom')
            if Dossiers.objects.filter(id_parent=dossier.id_parent,nom=nom).exclude(id_dossier=id_dossier).exists():
                return redirect("/archive/liste_dossier/" + str(user.id_user) + "/" + str(dossier.id_parent))
            else :
                dossier.nom = nom
                dossier.save()
                if dossier.id_parent :
                 return redirect("/archive/liste_dossier/" + str(user.id_user) + "/" + str(dossier.id_parent))
                else :
                    return redirect("/archive/dossier_racine/" + str(user.id_user))
        else :
            id_dossier = request.POST.get('id_dossier')
            nouv_direction = request.POST.get('nouv_direction')
            nom = request.POST.get('nom')
            dossier=Dossiers.objects.get(id_dossier=id_dossier)
            dossier.direction=nouv_direction
            dossier.nom=nom
            dossier.save()
            return redirect("/archive/dossier_racine/" + str(user.id_user))

################Gestion MnAGEMENT


def liste_tout_document(request, id_user):
    # Obtenir le dossier courant
    # Récupérer les informations de l'utilisateur
    user = Utilisateur.objects.get(id_user=id_user)
    tout_document=Documents.objects.all()
    tout_dossier = Dossiers.objects.all()
    return render(request, 'templatetra/tout_document_archive.html', {
        'util': user,
        'aff_tt' : 'tout',
        'tout_doc': tout_document,
        'tout_dos': tout_dossier,
         # Pour avoir l'ordre du plus ancien au plus récent
    })

######Fetch documents
from django.http import JsonResponse

def fetch_documents(request):
    documents = Documents.objects.all()
    data = []
    for doc in documents:
        data.append({
            'id_document': doc.id_document,
            'nom': doc.nom,
            'direction': doc.direction,
            'numero_document': doc.numero_document,
            'date_creation': doc.date_creation,
            'eta': doc.eta,
            'bl': doc.bl,
            'client': doc.client,
            'vessel': doc.vessel,
            'numero_voyage': doc.numero_voyage,
            'journal': doc.journal,
            'montant': doc.montant,
            'direction': doc.direction
        })
    return JsonResponse({'documents': data})




##############Edit Document
# views.py

from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from docx import Document
import os


def edit_document(request, id_user, id_doc):
    # Récupérer le document depuis la base de données
    doc = get_object_or_404(Documents, id_document=id_doc)
    user = get_object_or_404(Utilisateur, id_user=id_user)
    fs = FileSystemStorage()
    file_path = fs.path(
        os.path.join('C:\\MEDLOG\\gestion_archive\\inventory-management\\static\\archives\\', str(doc.nom)))

    # Ouvrir le document existant
    with open(file_path, 'rb') as file:
        document = Document(file)

    if request.method == 'POST':
        new_text = request.POST.get('content', '')

        # Mettre à jour le contenu du document avec le nouveau texte
        for para in document.paragraphs:
            para.text = new_text

        # Sauvegarder les modifications dans un objet BytesIO
        modified_file = BytesIO()
        document.save(modified_file)
        modified_file.seek(0)

        # Créer une réponse HTTP pour le téléchargement du document modifié
        response = HttpResponse(modified_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=modified_document.docx'
        return response

    # Convertir le document en texte pour l'affichage dans la zone de texte
    content = '\n'.join([para.text for para in document.paragraphs])
    return render(request, 'templatetra/edit_document.html', {'content': content, 'user': user, 'doc': doc})
def sauvegarde_contenu_document(request):
    if request.method == 'POST':
        id_user = request.POST.get('id_user')
        id_doc = request.POST.get('id_doc')
        content = request.POST.get('content', '')

        if not id_user or not id_doc:
            return redirect('error_page')  # Rediriger vers une page d'erreur si les identifiants sont manquants

        # Récupérer le document depuis la base de données
        doc = get_object_or_404(Documents, id_document=id_doc)
        fs = FileSystemStorage()
        file_path = fs.path(os.path.join('C:\\MEDLOG\\gestion_archive\\inventory-management\\static\\archives\\', str(doc.nom)))

        # Ouvrir le document existant
        with open(file_path, 'rb') as file:
            document = Document(file)

        # Mettre à jour le contenu du document avec le nouveau texte
        for para in document.paragraphs:
            para.text = content

        # Sauvegarder les modifications dans un objet BytesIO
        modified_file = BytesIO()
        document.save(modified_file)
        modified_file.seek(0)

        # Chemin pour enregistrer le document modifié
        modified_file_path = os.path.join('C:\\MEDLOG\\gestion_archive\\inventory-management\\static\\archives\\', 'modified_' + str(doc.nom))

        # Sauvegarder le document modifié
        with open(modified_file_path, 'wb') as file:
            file.write(modified_file.getvalue())

        # Rediriger vers la page de modification ou de téléchargement
        return redirect('edit_document', id_user=id_user, id_doc=id_doc)

    # Rediriger vers une page d'erreur si la requête n'est pas POST
    return redirect('error_page')

#####Udate Document

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
from .models import Documents  # Assurez-vous que ce modèle existe et est configuré correctement


def update_document(request):
    if request.method == 'POST':
        # Récupération des données du formulaire
        id_user = request.POST.get('id_user')
        doscrt = request.POST.get('doscrt')
        document_id = request.POST.get('document_id')
        nom = request.POST.get('nom')
        numero_document = request.POST.get('numero_document')
       # date_creation_str = request.POST.get('date_creation')
        eta_str = request.POST.get('eta') or '1900-01-01T00:00'
        bl = request.POST.get('bl')
        client = request.POST.get('client')
        vessel = request.POST.get('vessel')
        numero_voyage = request.POST.get('numero_voyage')
        journal = request.POST.get('journal')
        montant = request.POST.get('montant')
        if request.POST.get('detailtt') == 'detailtt':

            ###########3Gestion fichier
            try:
                eta = datetime.strptime(eta_str, '%Y-%m-%dT%H:%M')
            except ValueError:
                return HttpResponse("Format de ETA invalide. Le format attendu est 'YYYY-MM-DDTHH:MM'.", status=400)

            # Récupérer le document à mettre à jour
            try:
                document = Documents.objects.get(id_document=document_id)
            except Documents.DoesNotExist:
                return HttpResponse("Document non trouvé", status=404)

            # Mettre à jour les propriétés du document
            fichier = request.FILES.get('fichier')
            if fichier:
                file_path = 'C:\\MEDLOG\\gestion_archive\\inventory-management\\static\\archives\\' + str(nom)
                if os.path.exists(file_path):

                    os.remove(file_path)
                    handle_uploaded_file(fichier)
                    current_timestamp = int(datetime.timestamp(datetime.now()))
                    new_nom = str(current_timestamp) + fichier.name
                    document.nom = new_nom
            else :
             document.nom = nom
            document.numero_document = numero_document
           # document.date_creation = date_creation
            document.eta = eta
            document.bl = bl
            document.client = client
            document.vessel = vessel
            document.numero_voyage = numero_voyage
            document.journal = journal
            document.montant = montant

            # Enregistrer les modifications
            document.save()

            return redirect(f'/archive/liste_document_direction/{id_user}/{doscrt}')
        else :
            ###########3Gestion fichier
            try:
                eta = datetime.strptime(eta_str, '%Y-%m-%dT%H:%M')
            except ValueError:
                return HttpResponse("Format de ETA invalide. Le format attendu est 'YYYY-MM-DDTHH:MM'.", status=400)

            # Récupérer le document à mettre à jour
            try:
                document = Documents.objects.get(id_document=document_id)
            except Documents.DoesNotExist:
                return HttpResponse("Document non trouvé", status=404)

            # Mettre à jour les propriétés du document
            fichier = request.FILES.get('fichier')
            if fichier:
                file_path = 'C:\\MEDLOG\\gestion_archive\\inventory-management\\static\\archives\\' + str(nom)
                if os.path.exists(file_path):
                    os.remove(file_path)
                    handle_uploaded_file(fichier)
                    current_timestamp = int(datetime.timestamp(datetime.now()))
                    new_nom = str(current_timestamp) + fichier.name
                    document.nom = new_nom
            else:
                document.nom = nom
            document.numero_document = numero_document
            # document.date_creation = date_creation
            document.eta = eta
            document.bl = bl
            document.client = client
            document.vessel = vessel
            document.numero_voyage = numero_voyage
            document.journal = journal
            document.montant = montant

            # Enregistrer les modifications
            document.save()

            return redirect(f'/archive/liste_dossier/{id_user}/{doscrt}')
            a=t

        return HttpResponse("Méthode non autorisée", status=405)


  # Assurez-vous que ce modèle existe et est configuré correctement

#################################  AJAX ###############################################
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Dossiers, Documents, Acces, Demandes

def get_descendants_data(request):
    dossier_crt = get_object_or_404(Dossiers, id_dossier=127)
    id_user = request.user.id

    def get_descendants(dossier):
        sous_dossiers = Dossiers.objects.filter(id_parent=dossier.id_dossier, direction=dossier.direction)
        documents = Documents.objects.filter(id_parent=dossier.id_dossier)
        all_sous_dossiers = list(sous_dossiers)
        all_documents = list(documents)

        for sous_dossier in sous_dossiers:
            desc_docs, desc_dossiers = get_descendants(sous_dossier)
            all_documents.extend(desc_docs)
            all_sous_dossiers.extend(desc_dossiers)

        return all_documents, all_sous_dossiers

    lst_document, lst_dossier = get_descendants(dossier_crt)
    lst_document = [doc for doc in lst_document if doc.id_parent == dossier_crt.id_dossier]

    def update_acces_and_etat(items, is_document=True):
        for item in items:
            try:
                acces = Acces.objects.get(id_utilisateur=id_user, etat=1, permission='acces',
                                          id_objet=(item.id_document if is_document else item.id_dossier))
                item.acces = acces.etat
            except Acces.DoesNotExist:
                item.acces = 5

            try:
                dmd = Demandes.objects.get(id_demandeur=id_user,
                                           id_objet=(item.id_document if is_document else item.id_dossier))
                item.etat = dmd.etat
            except Demandes.DoesNotExist:
                item.etat = 5

    update_acces_and_etat(lst_document)
    update_acces_and_etat(lst_dossier, is_document=False)

    arborescence = []
    get_parent = dossier_crt
    while get_parent:
        arborescence.append(get_parent)
        get_parent = Dossiers.objects.filter(id_dossier=get_parent.id_parent).first()

    tout_document = Documents.objects.all()
    for doc in tout_document:
        dossier_acces = Acces.objects.filter(etat=1, id_utilisateur=id_user, permission='acces', type='dossier')
        restricted_doc_ids = set(dossier_acces.values_list('id_objet', flat=True))
        doc.acces_dos = 1 if doc.id_parent in restricted_doc_ids else 0

    update_acces_and_etat(tout_document)

    # Convertir les listes d'objets en dictionnaires
    documents_data = [{'id_document': doc.id_document,'direction': doc.direction,  'nom': doc.nom, 'acces': doc.acces, 'etat': doc.etat} for doc in lst_document]
    dossiers_data = [{'id_dossier': dos.id_dossier,'direction': dos.direction,  'nom': dos.nom, 'acces': dos.acces, 'etat': dos.etat} for dos in lst_dossier]
    arborescence_data = [dossier.id_dossier for dossier in reversed(arborescence)]

    response_data = {
        'documents': documents_data,
        'dossiers': dossiers_data,
        'arborescence': arborescence_data
    }

    return JsonResponse(response_data)



##############fetch Dossier
# views.py

from django.http import JsonResponse
from .models import Dossiers, Utilisateur, Acces


def fetch_dossiers(request):
    id_user = request.GET.get('id_user')
    id_prt = request.GET.get('id_prt')

    if not id_user or not id_prt:
        return JsonResponse({'error': 'Missing parameters'}, status=400)

    try:
        dossier_crt = Dossiers.objects.get(id_dossier=id_prt)
        user = Utilisateur.objects.get(id_user=id_user)
    except (Dossiers.DoesNotExist, Utilisateur.DoesNotExist):
        return JsonResponse({'error': 'Invalid parameters'}, status=400)

    def get_descendants(dossier):
        sous_dossiers = Dossiers.objects.filter(id_parent=dossier.id_dossier, direction=user.direction)
        all_sous_dossiers = list(sous_dossiers)
        for sous_dossier in sous_dossiers:
            desc_dossiers = get_descendants(sous_dossier)
            all_sous_dossiers.extend(desc_dossiers)
        return all_sous_dossiers

    lst_dossier = get_descendants(dossier_crt)

    def update_acces_and_etat(items):
        for item in items:
            try:
                acces = Acces.objects.get(id_utilisateur=id_user, etat=1, permission='acces', id_objet=item.id_dossier)
                item.acces = acces.etat
            except Acces.DoesNotExist:
                item.acces = 5

    update_acces_and_etat(lst_dossier)

    dossiers_data = [
        {'id_dossier': d.id_dossier,'direction': d.direction, 'nom': d.nom, 'date_creation': d.date_creation,
         'id_parent': d.id_parent, 'acces': d.acces} for d in lst_dossier]

    return JsonResponse({'lst_dos': dossiers_data})









##################Fetch document et Dossier
# views.py

from django.http import JsonResponse
from .models import Dossiers, Utilisateur, Acces, Documents, Demandes
from django.http import JsonResponse


def fetch_dossiers_tout(request):
    id_user = request.GET.get('id_user')
    id_prt = request.GET.get('id_prt')

    if not id_user or not id_prt:
        return JsonResponse({'error': 'Missing parameters'}, status=400)

    try:
        dossier_crt = Dossiers.objects.get(id_dossier=id_prt)
        user = Utilisateur.objects.get(id_user=id_user)
    except (Dossiers.DoesNotExist, Utilisateur.DoesNotExist):
        return JsonResponse({'error': 'Invalid parameters'}, status=400)



    lst_document=Documents.objects.filter(id_parent=id_prt)
    #if user.direction == 'management' :
        #dos_tom=Dossiers.objects.get(id_parent__isnull=True)
        #lst_dossier = Dossiers.objects.filter(id_parent=dos_tom.id_dossier)
    #else :
    lst_dossier = Dossiers.objects.filter(id_parent=id_prt)
    def update_acces_and_etat(items, is_document=True):
        for item in items:
            try:
                acces = Acces.objects.get(id_utilisateur=id_user, etat=1, permission='acces',
                                          id_objet=(item.id_document if is_document else item.id_dossier))
                item.acces = acces.etat
            except Acces.DoesNotExist:
                item.acces = 5

            try:
                dmd = Demandes.objects.get(id_demandeur=id_user,
                                           id_objet=(item.id_document if is_document else item.id_dossier))
                item.etat = dmd.etat
            except Demandes.DoesNotExist:
                item.etat = 5
            ##########Appartient Mnagement
            try:
                dmd = Demandes.objects.get(id_demandeur=id_user,
                                           id_objet=(item.id_document if is_document else item.id_dossier))
                item.etat = dmd.etat
            except Demandes.DoesNotExist:
                item.etat = 5

    update_acces_and_etat(lst_document)
    update_acces_and_etat(lst_dossier, is_document=False)
    # Gestion de tous les documents
    tout_document = Documents.objects.all()
    update_acces_and_etat(tout_document)

    dossiers_data = [
        {'id_dossier': d.id_dossier, 'nom': d.nom, 'direction': d.direction, 'date_creation': d.date_creation,
         'id_parent': d.id_parent, 'acces': d.acces} for d in lst_dossier]

    documents_data = [
        {'id_document': d.id_document,'direction': d.direction,'nom': d.nom, 'numero_document': d.numero_document, 'date_creation': d.date_creation,
         'eta': d.eta, 'bl': d.bl, 'client': d.client, 'vessel': d.vessel, 'numero_voyage': d.numero_voyage,
         'journal': d.journal, 'montant': d.montant, 'acces': d.acces, 'etat': d.etat}
        for d in lst_document]
    tout_documents_data = [
        {'id_document': d.id_document,'direction': d.direction,'nom': d.nom, 'numero_document': d.numero_document,
         'date_creation': d.date_creation,
         'eta': d.eta, 'bl': d.bl, 'client': d.client, 'vessel': d.vessel, 'numero_voyage': d.numero_voyage,
         'journal': d.journal, 'montant': d.montant, 'acces': d.acces, 'etat': d.etat}
        for d in tout_document]

    return JsonResponse({'lst_dos': dossiers_data, 'lst_docs': documents_data,'tout_docs': tout_documents_data})



####Fetch Documnets Direction
@csrf_exempt
def  liste_tout_direction(request, id_user):
    user = Utilisateur.objects.get(id_user=id_user)
    return render(request, 'templatetra/tout_document_direction.html', { 'user': user, 'util': user})
def fetch_dossiers_tout_direction(request):
    id_user = request.GET.get('id_user')

    def update_acces_and_etat(items, is_document=True):
        for item in items:
            try:
                acces = Acces.objects.get(id_utilisateur=id_user, etat=1, permission='acces',
                                          id_objet=(item.id_document if is_document else item.id_dossier))
                item.acces = acces.etat
            except Acces.DoesNotExist:
                item.acces = 5

            try:
                dmd = Demandes.objects.get(id_demandeur=id_user,
                                           id_objet=(item.id_document if is_document else item.id_dossier))
                item.etat = dmd.etat
            except Demandes.DoesNotExist:
                item.etat = 5


    # Gestion de tous les documents
    user=Utilisateur.objects.get(id_user=id_user)
    dossier_dir=Dossiers.objects.filter(direction=user.direction)
    restricted_doc_ids = set(dossier_dir.values_list('id_parent', flat=True))
    tout_document = Documents.objects.filter(id_parent__in=restricted_doc_ids)
    update_acces_and_etat(tout_document)


    tout_documents_data = [
        {'id_document': d.id_document, 'nom': d.nom, 'direction': d.direction, 'numero_document': d.numero_document,
         'date_creation': d.date_creation,
         'eta': d.eta, 'bl': d.bl, 'client': d.client, 'vessel': d.vessel, 'numero_voyage': d.numero_voyage,
         'journal': d.journal, 'montant': d.montant, 'acces': d.acces, 'etat': d.etat}
        for d in tout_document]

    return JsonResponse({ 'lst_docs':  tout_documents_data})





##############
def restraindre_page_tout(request, id_user, id_objet, type):
    dossier_crt = False
    user = Utilisateur.objects.get(id_user=id_user)

    # user_direction = Utilisateur.objects.filter(direction=user.direction).exclude(user)
    acces_users = Acces.objects.filter(etat=1, id_objet=id_objet, type=type).values_list('id_utilisateur', flat=True)

    #####Agent direction
    agents_dir = Utilisateur.objects.filter(direction=user.direction).exclude(id_user__in=acces_users).exclude(
        id_user=user.id_user)
    ###################Agent No direction
    agents_no_dir = Utilisateur.objects.filter(id_user__in=acces_users).exclude(direction=user.direction)
    if type == 'dossier':
        objet = Dossiers.objects.get(id_dossier=id_objet)
        # id_objet = objet.id_dossier
        dossier_crt = Dossiers.objects.filter(id_dossier=objet.id_parent).first()
        if not dossier_crt:
            dossier_crt = Dossiers.objects.filter(id_dossier=objet.id_dossier).first()
        id_dos_crt = dossier_crt.id_dossier
        type = 'dossier'
    else:
        objet = Documents.objects.get(id_document=id_objet)
        # id_objet=objet.id_document
        dossier_crt = Dossiers.objects.get(id_dossier=objet.id_parent)
        id_dos_crt = dossier_crt.id_dossier
        type = 'document'

    return render(request, 'templatetra/restreindre_tout_doc_dirrection.html', {
        'util': user,
        'agents_no_dir': agents_no_dir,
        'agents_dir': agents_dir,
        'type': type,
        'id_objet': id_objet,
        'id_dos_crt': id_dos_crt  # Pour avoir l'ordre du plus ancien au plus récent
    })


#########Affiche tout

def affiche_document_tout(request, id_user, id_doc):
    document = Documents.objects.get(id_document=id_doc)
    dossier_crt = Dossiers.objects.get(id_dossier=document.id_parent)
    user = Utilisateur.objects.get(id_user=id_user)
    base_dir = r'C:\MEDLOG\gestion_archive\inventory-management\static\archives'
    file_path = os.path.join(base_dir, document.nom)
    file_extension = os.path.splitext(file_path)[1].lower()
    try:
        if file_extension == '.pdf':
            return render(request, 'templatetra/affiche_document_tout.html', {
                'util': user,
                #'aff_tt':'aff',
                'doc': document,
                'dossier_crt': dossier_crt,
                'file_url': f"/static/archives/{document.nom}",
                'file_type': 'pdf'
            })
        elif file_extension in ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx']:
            output_file = convert_to_pdf(file_path)
            return render(request, 'templatetra/affiche_document_tout.html', {
                'util': user,
                'doc': document,
                'dossier_crt': dossier_crt,
                'file_url': f"/static/archives/{os.path.basename(output_file)}",
                'file_type': 'pdf'
            })
        else:
            return HttpResponse("Type de fichier non pris en charge.", status=415)
    except Exception as e:
        response = HttpResponse(open(file_path, 'rb').read(), content_type='application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
        return response






def redirige_app(request,app):
    if app == 'archibve' :
       return render(request, 'templatetra/login1.html')
    else :
        return render(request, 'templatetra/login_ged.html')

    ################################ FONCTION QUI RETOURNE LA PAGE DE MODIFICATION DU MOT D EPASSE
def profile_page (request, id_user):
    util = Utilisateur.objects.get(id_user=id_user)
    return render(request, 'templatetra/page_profile.html', {'util':util,})


def modifier_profile(request):
    message = ''

    if request.method == 'POST':
        id_user = request.POST.get('id_user')

        try:
            user = Utilisateur.objects.get(id_user=id_user)
        except Utilisateur.DoesNotExist:
            message = 'Utilisateur introuvable.'
            return render(request, 'templatetra/page_profile.html', {'message': message})

        fullname = request.POST.get('fullname', user.fullname)  # Valeur par défaut si non fournie
        email = request.POST.get('email', user.email)

        # Vérifier si l'email est déjà utilisé par un autre utilisateur
        if Utilisateur.objects.filter(email=email).exclude(id_user=id_user).exists():
            message = 'Cet email est déjà utilisé !!'
            return render(request, 'templatetra/page_profile.html', {'util': user, 'message': message})

        password = request.POST.get('password')

        # Mettre à jour les informations de l'utilisateur
        user.fullname = fullname
        user.email = email

        if password:
            user.password=password  # Utiliser set_password pour hacher le mot de passe

        user.save()

        messages.success(request, "Profil mis à jour avec succès.")
        return render(request, 'templatetra/index.html', {'util': user, 'message': message})




def retour_profile (request, id_user):
    util = Utilisateur.objects.get(id_user=id_user)
    return render(request, 'templatetra/index.html', {'util':util,})
